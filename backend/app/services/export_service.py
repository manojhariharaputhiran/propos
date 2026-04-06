"""
Export Service: generates Word, PDF, PPT, and Excel exports for a completed RFP.
"""
import io
import os
import uuid
import tempfile
from typing import List, Dict, Any, Optional

from app.agents.design_agent import build_ppt


# ── Word (python-docx) ────────────────────────────────────────────────────────

def export_word(
    title: str,
    questions: List[Dict[str, Any]],
    critic_report: Optional[Dict[str, Any]] = None,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Critic scorecard
    if critic_report:
        doc.add_heading("Proposal Scorecard", level=1)
        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        headers = ["Win Probability", "Completeness", "Persuasiveness", "Compliance"]
        keys = ["win_probability", "completeness", "persuasiveness", "compliance"]
        for i, (h, k) in enumerate(zip(headers, keys)):
            table.cell(0, i).text = h
            table.cell(1, i).text = f"{critic_report.get(k, 0)}%"

        if critic_report.get("suggestions"):
            doc.add_heading("Suggested Improvements", level=2)
            for s in critic_report["suggestions"]:
                doc.add_paragraph(s, style="List Bullet")

        if critic_report.get("critical_flags"):
            doc.add_heading("Critical Flags", level=2)
            for f in critic_report["critical_flags"]:
                p = doc.add_paragraph(f, style="List Bullet")
                p.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        doc.add_page_break()

    # Q&A
    doc.add_heading("Proposal Responses", level=1)
    for i, q in enumerate(questions):
        doc.add_heading(f"Q{i+1}: {q['text']}", level=2)
        answer = q.get("final_answer") or q.get("draft_answer") or q.get("draft") or "Pending."
        p = doc.add_paragraph(answer)

        conf = q.get("confidence")
        if conf is not None:
            conf_para = doc.add_paragraph(f"Confidence: {int(conf * 100)}%")
            conf_para.runs[0].font.size = Pt(9)
            conf_para.runs[0].font.italic = True

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF (reportlab fallback) ──────────────────────────────────────────────────

def export_pdf(
    title: str,
    questions: List[Dict[str, Any]],
    critic_report: Optional[Dict[str, Any]] = None,
) -> bytes:
    """
    Generate PDF. First tries to convert Word doc via libreoffice;
    falls back to reportlab if libreoffice not available.
    """
    try:
        return _pdf_via_libreoffice(title, questions, critic_report)
    except Exception:
        return _pdf_via_reportlab(title, questions, critic_report)


def _pdf_via_libreoffice(title, questions, critic_report) -> bytes:
    import subprocess, shutil

    if not shutil.which("soffice") and not shutil.which("libreoffice"):
        raise RuntimeError("libreoffice not found")

    docx_bytes = export_word(title, questions, critic_report)
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "rfp.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path]
        subprocess.run(cmd, check=True, capture_output=True, timeout=30)

        pdf_path = os.path.join(tmpdir, "rfp.pdf")
        with open(pdf_path, "rb") as f:
            return f.read()


def _pdf_via_reportlab(title, questions, critic_report) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()

    primary = HexColor("#1E40AF")
    story = []

    # Title
    title_style = ParagraphStyle("Title", parent=styles["Title"], textColor=primary, fontSize=24)
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.3 * inch))

    # Scorecard
    if critic_report:
        story.append(Paragraph("Proposal Scorecard", styles["Heading1"]))
        data = [
            ["Win Probability", "Completeness", "Persuasiveness", "Compliance"],
            [
                f"{critic_report.get('win_probability', 0)}%",
                f"{critic_report.get('completeness', 0)}%",
                f"{critic_report.get('persuasiveness', 0)}%",
                f"{critic_report.get('compliance', 0)}%",
            ],
        ]
        t = Table(data, colWidths=[1.5 * inch] * 4)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), primary),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2 * inch))

    # Q&A
    story.append(Paragraph("Proposal Responses", styles["Heading1"]))
    for i, q in enumerate(questions):
        story.append(Paragraph(f"Q{i+1}: {q['text']}", styles["Heading2"]))
        answer = q.get("final_answer") or q.get("draft_answer") or q.get("draft") or "Pending."
        story.append(Paragraph(answer, styles["BodyText"]))
        story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    return buf.getvalue()


# ── Excel (openpyxl) ──────────────────────────────────────────────────────────

def export_sheets(
    title: str,
    questions: List[Dict[str, Any]],
    critic_report: Optional[Dict[str, Any]] = None,
) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # ── Q&A Sheet ──────────────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Proposal Q&A"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1E40AF")
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers = ["#", "Question", "Draft Answer", "Confidence", "Needs SME", "Final Answer"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for i, q in enumerate(questions, 2):
        ws.cell(row=i, column=1, value=i - 1)
        ws.cell(row=i, column=2, value=q.get("text", ""))
        ws.cell(row=i, column=3, value=q.get("draft_answer") or q.get("draft") or "")
        conf = q.get("confidence")
        ws.cell(row=i, column=4, value=f"{int(conf * 100)}%" if conf is not None else "")
        ws.cell(row=i, column=5, value="Yes" if q.get("needs_sme") else "No")
        ws.cell(row=i, column=6, value=q.get("final_answer") or "")

        for col in range(1, 7):
            ws.cell(row=i, column=col).border = thin
            ws.cell(row=i, column=col).alignment = Alignment(wrap_text=True, vertical="top")

    # Column widths
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 50
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 50

    # ── Scorecard Sheet ────────────────────────────────────────────────────────
    if critic_report:
        ws2 = wb.create_sheet("Scorecard")
        ws2.append(["Metric", "Score"])
        ws2["A1"].font = header_font
        ws2["A1"].fill = header_fill
        ws2["B1"].font = header_font
        ws2["B1"].fill = header_fill

        for key in ["win_probability", "completeness", "persuasiveness", "compliance"]:
            ws2.append([key.replace("_", " ").title(), f"{critic_report.get(key, 0)}%"])

        ws2.append([])
        ws2.append(["Suggestions"])
        for s in critic_report.get("suggestions", []):
            ws2.append(["", s])

        ws2.column_dimensions["A"].width = 20
        ws2.column_dimensions["B"].width = 60

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PPT wrapper ────────────────────────────────────────────────────────────────

def export_ppt(
    title: str,
    questions: List[Dict[str, Any]],
    critic_report: Optional[Dict[str, Any]] = None,
    template_path: Optional[str] = None,
) -> bytes:
    return build_ppt(
        questions=questions,
        title=title,
        critic_report=critic_report,
        template_path=template_path,
    )
